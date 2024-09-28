from docx import Document

# Create a new Document
doc = Document()

# Title Slide
doc.add_heading('Hikvision NVR Configuration Guide', 0)
doc.add_paragraph('Using the SADP Tool to Discover NVR and IP Camera IP Addresses\n')
doc.add_paragraph('Presented by: Your Name/Company\n')

# Introduction Slide
doc.add_heading('Introduction', level=1)
doc.add_paragraph(
    "Hikvision NVRs are widely used in video surveillance systems due to their reliability and compatibility with IP cameras. "
    "In this guide, we will cover the complete process of configuring a Hikvision NVR, from finding the NVR and IP camera IP addresses "
    "to setting up the device through its web interface. Understanding how to configure these settings is essential for optimal performance and monitoring.\n"
    "Steps included:\n"
    "- Using the SADP Tool to discover IP addresses\n"
    "- Modifying IP settings to match your network\n"
    "- Logging into the NVR and configuring the system."
)

# SADP Tool Overview Slide
doc.add_heading('What is the SADP Tool?', level=1)
doc.add_paragraph(
    "The SADP (Search Active Devices Protocol) Tool is a utility provided by Hikvision to automatically detect all Hikvision devices, such as NVRs, DVRs, and IP cameras, "
    "on the same network. It simplifies the configuration process by allowing you to find devices, modify their IP addresses, and update firmware, among other features.\n\n"
    "Key Features:\n"
    "- Scans the local network to discover devices\n"
    "- Allows modification of IP addresses directly\n"
    "- Displays device information such as model, port, and status\n\n"
    "This tool is essential for setting up new devices and troubleshooting network issues."
)

# Downloading & Installing SADP Tool
doc.add_heading('How to Download and Install the SADP Tool', level=1)
doc.add_paragraph(
    "To get started, you need to download the SADP Tool from the official Hikvision website:\n\n"
    "1. Visit the Hikvision website: https://www.hikvision.com/en/support/tools/\n"
    "2. Locate the SADP Tool and download the appropriate version for your operating system (Windows or macOS).\n"
    "3. Run the installer and follow the on-screen instructions to complete the installation.\n\n"
    "After installation, you can launch the tool to begin scanning your network for Hikvision devices."
)

# Finding Devices with SADP Tool
doc.add_heading('Finding NVR and Camera IP Addresses Using SADP', level=1)
doc.add_paragraph(
    "Once you have installed the SADP Tool, follow these steps to find your Hikvision NVR and camera IP addresses:\n\n"
    "1. Open the SADP Tool; it will automatically start scanning your network for connected devices.\n"
    "2. Wait for the list of devices to populate. You will see columns for IP Address, Device Type, Port, and Status.\n"
    "3. Locate your NVR and cameras by identifying their respective Device Types (e.g., NVR, IP Camera).\n\n"
    "Screenshot Example: This screen will show all available devices, including their current IP addresses.\n"
    "Explanation of key fields:\n"
    "- **IP Address**: The current network address of the device.\n"
    "- **Device Type**: Type of device (e.g., NVR, camera).\n"
    "- **Port**: The port number used for accessing the device."
)

# Modifying IP Address
doc.add_heading('Changing IP Addresses with SADP Tool', level=1)
doc.add_paragraph(
    "If the NVR or camera's IP address does not match your network range, you can modify it using the SADP Tool:\n\n"
    "1. Select the NVR or camera from the device list.\n"
    "2. Enter the new IP address in the IP Address field. Ensure that the new IP falls within your network range (e.g., 192.168.1.x).\n"
    "3. Enter the device’s password to apply the changes.\n"
    "4. Click the ‘Modify’ button. The device will restart and apply the new IP configuration.\n\n"
    "Screenshot: Showing the process of entering the new IP address in SADP."
)

# Logging into the NVR Web Interface
doc.add_heading('Logging into the NVR', level=1)
doc.add_paragraph(
    "After modifying the IP address, you can access the NVR's web interface to configure it:\n\n"
    "1. Open a web browser and enter the NVR’s new IP address in the address bar (e.g., http://192.168.1.100).\n"
    "2. You will be prompted to log in. By default, Hikvision NVRs use the following credentials:\n"
    "   - **Username**: admin\n"
    "   - **Password**: admin123 (or the password set during initial setup)\n"
    "3. After logging in, you will be taken to the NVR’s dashboard, where you can start the configuration process.\n\n"
    "Screenshot: NVR login page with fields for Username and Password."
)

# Configuring the NVR
doc.add_heading('Configuring Basic NVR Settings', level=1)
doc.add_paragraph(
    "Once logged in, you can configure the essential settings of the NVR:\n\n"
    "1. **Date and Time**: Go to System > General to set the correct date and time.\n"
    "2. **Network Settings**: Under Network, set the NVR to use a static IP address if necessary.\n"
    "3. **Camera Setup**: Navigate to Camera Management and add IP cameras either manually or automatically.\n\n"
    "Step-by-step screenshots of these configuration processes are recommended."
)

# Adding IP Cameras to the NVR
doc.add_heading('Adding IP Cameras to the NVR', level=1)
doc.add_paragraph(
    "To add cameras to the NVR:\n\n"
    "1. Go to **Camera Management** in the NVR interface.\n"
    "2. Click ‘Add Camera’. If your cameras are in the same network range, they may be discovered automatically.\n"
    "3. If adding manually, enter the camera’s IP address, port, username, and password.\n"
    "4. Save the settings and check the live feed to ensure the camera is properly connected.\n\n"
    "Screenshot: Showing the Camera Management page with options to add cameras."
)

# Troubleshooting
doc.add_heading('Common Issues & Troubleshooting', level=1)
doc.add_paragraph(
    "Here are common issues you may encounter during setup and how to solve them:\n\n"
    "- **Device not found in SADP**: Ensure that the device is powered on and connected to the same network.\n"
    "- **Incorrect IP range**: If the device has an IP address outside your network range, temporarily set your PC to match the device’s subnet.\n"
    "- **Unable to log into the NVR**: Double-check the default login credentials and ensure there is no caps lock or input error in the password.\n"
    "- **Camera not connecting to the NVR**: Check the camera’s IP address, verify the correct credentials, and confirm network settings."
)

# Conclusion Slide
doc.add_heading('Conclusion & Best Practices', level=1)
doc.add_paragraph(
    "Setting up a Hikvision NVR involves several key steps, from discovering devices with the SADP Tool to configuring the NVR through its web interface. "
    "Once set up, regular monitoring and firmware updates are essential for maintaining optimal system performance.\n\n"
    "Best practices include:\n"
    "- Always use static IP addresses for the NVR and cameras.\n"
    "- Regularly back up configuration files and firmware.\n"
    "- Monitor network activity to ensure devices stay online and accessible."
)

# Q&A Slide
doc.add_heading('Questions & Answers', level=1)
doc.add_paragraph("Feel free to ask any questions about the configuration process or troubleshooting steps.")

# Save the detailed Document
file_path_detailed = "Hikvision_NVR_Configuration_Guide_Detailed.docx"
doc.save(file_path_detailed)

print(f"Document saved as {file_path_detailed}")
